"""
Loader per file Word (.docx).

Un file .docx è internamente un archivio ZIP contenente file XML.
Il testo è frammentato in "run" — segmenti con la stessa formattazione.
python-docx li riassembla in paragrafi.

Gestione delle tabelle: i documenti Word spesso contengono tabelle
(contratti con clausole, reportistica). Le estraiamo riga per riga
nello stesso formato dell'ExcelLoader per consistenza.
"""

from pathlib import Path
from docx import Document as DocxDocument

from rag_assistant.adapters.base import DocumentLoader
from rag_assistant.core.models import Document


class WordLoader(DocumentLoader):
    """Carica file Word (.docx) estraendo testo e tabelle."""

    def load(self, file_path: str) -> list[Document]:
        """Carica un file Word.

        Estrae sia il testo dei paragrafi che il contenuto delle tabelle.
        Le tabelle vengono convertite nello stesso formato header:valore
        usato dall'ExcelLoader.

        Returns:
            Lista con un singolo Document.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File non trovato: {file_path}")

        docx = DocxDocument(str(path))
        parts = []

        for element in docx.element.body:
            tag = element.tag.split("}")[-1]  # Rimuove il namespace XML

            if tag == "p":
                # È un paragrafo
                para = self._extract_paragraph_text(element, docx)
                if para.strip():
                    parts.append(para)

            elif tag == "tbl":
                # È una tabella
                table_text = self._extract_table(element, docx)
                if table_text.strip():
                    parts.append(table_text)

        full_text = "\n\n".join(parts)

        return [Document(
            source_path=str(path.resolve()),
            source_name=path.name,
            doc_type="docx",
            text=full_text,
            metadata={
                "paragraphs": len(docx.paragraphs),
                "tables": len(docx.tables),
            },
        )]

    def _extract_paragraph_text(self, element, docx) -> str:
        """Estrae il testo da un elemento paragrafo XML."""
        # Cerca tutti i 'run' (segmenti di testo) dentro il paragrafo
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        runs = element.findall(f".//{namespace}t")
        return "".join(run.text or "" for run in runs)

    def _extract_table(self, element, docx) -> str:
        """Converte una tabella Word in testo leggibile.

        Stessa logica dell'ExcelLoader: prima riga = header,
        righe successive = "Header: Valore | Header: Valore".
        """
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        rows = element.findall(f".//{namespace}tr")

        if not rows:
            return ""

        # Funzione helper per estrarre il testo da una cella
        def cell_text(cell_element) -> str:
            texts = cell_element.findall(f".//{namespace}t")
            return " ".join(t.text or "" for t in texts).strip()

        # Prima riga = headers
        header_cells = rows[0].findall(f".//{namespace}tc")
        headers = [cell_text(c) or f"Col_{i}" for i, c in enumerate(header_cells)]

        # Righe successive = dati
        lines = []
        for row in rows[1:]:
            cells = row.findall(f".//{namespace}tc")
            parts = []
            for j, cell in enumerate(cells):
                value = cell_text(cell)
                if not value:
                    continue
                header = headers[j] if j < len(headers) else f"Col_{j}"
                parts.append(f"{header}: {value}")

            if parts:
                lines.append(" | ".join(parts))

        return "\n".join(lines)

    @staticmethod
    def supported_extensions() -> list[str]:
        return [".docx"]
