"""Test per i document loader.

Testa ogni loader con file creati in memoria o temporanei.
Non testiamo file reali della campagna (quelli sono dati privati),
ma file sintetici che coprono i casi principali.
"""

import csv
import tempfile
from pathlib import Path

import openpyxl
import pytest
from docx import Document as DocxDocument

from rag_assistant.adapters.txt_loader import TXTLoader
from rag_assistant.adapters.pdf_loader import PDFLoader
from rag_assistant.adapters.excel_loader import ExcelLoader
from rag_assistant.adapters.word_loader import WordLoader
from rag_assistant.adapters.csv_loader import CSVLoader
from rag_assistant.adapters.loader_registry import get_loader, supported_extensions


class TestTXTLoader:
    def test_load_simple(self, tmp_path):
        """Carica un file di testo semplice."""
        file = tmp_path / "test.txt"
        file.write_text("Riga uno\nRiga due\nRiga tre", encoding="utf-8")

        loader = TXTLoader()
        docs = loader.load(str(file))

        assert len(docs) == 1
        assert docs[0].doc_type == "txt"
        assert "Riga uno" in docs[0].text
        assert "Riga tre" in docs[0].text

    def test_file_not_found(self):
        loader = TXTLoader()
        with pytest.raises(FileNotFoundError):
            loader.load("/non/esiste.txt")

    def test_supported_extensions(self):
        assert TXTLoader.supported_extensions() == [".txt"]


class TestPDFLoader:
    def test_supported_extensions(self):
        assert PDFLoader.supported_extensions() == [".pdf"]

    def test_file_not_found(self):
        loader = PDFLoader()
        with pytest.raises(FileNotFoundError):
            loader.load("/non/esiste.pdf")


class TestExcelLoader:
    def test_load_single_sheet(self, tmp_path):
        """Carica un Excel con un foglio e verifica la conversione."""
        file = tmp_path / "test.xlsx"

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Fornitori"
        ws.append(["Fornitore", "Prodotto", "Prezzo"])
        ws.append(["Ferrara", "Uva Italia", 1.20])
        ws.append(["Rossi", "Pomodoro", 0.80])
        wb.save(str(file))

        loader = ExcelLoader()
        docs = loader.load(str(file))

        assert len(docs) == 1
        assert docs[0].doc_type == "excel"
        assert "Fornitori" in docs[0].source_name

        # Verifica che il testo contenga header:valore
        assert "Fornitore: Ferrara" in docs[0].text
        assert "Prodotto: Uva Italia" in docs[0].text
        assert "Fornitore: Rossi" in docs[0].text

    def test_load_multi_sheet(self, tmp_path):
        """Un Excel con 2 fogli produce 2 Document."""
        file = tmp_path / "multi.xlsx"

        wb = openpyxl.Workbook()
        ws1 = wb.active
        ws1.title = "Vendite"
        ws1.append(["Mese", "Totale"])
        ws1.append(["Gennaio", 15000])

        ws2 = wb.create_sheet("Acquisti")
        ws2.append(["Fornitore", "Importo"])
        ws2.append(["Ferrara", 5000])
        wb.save(str(file))

        loader = ExcelLoader()
        docs = loader.load(str(file))

        assert len(docs) == 2
        names = [d.source_name for d in docs]
        assert any("Vendite" in n for n in names)
        assert any("Acquisti" in n for n in names)

    def test_empty_sheet_skipped(self, tmp_path):
        """Un foglio vuoto non produce un Document."""
        file = tmp_path / "empty.xlsx"

        wb = openpyxl.Workbook()
        ws1 = wb.active
        ws1.title = "Dati"
        ws1.append(["Col1", "Col2"])
        ws1.append(["A", "B"])

        wb.create_sheet("Vuoto")  # Foglio vuoto
        wb.save(str(file))

        loader = ExcelLoader()
        docs = loader.load(str(file))

        assert len(docs) == 1
        assert "Dati" in docs[0].source_name

    def test_file_not_found(self):
        loader = ExcelLoader()
        with pytest.raises(FileNotFoundError):
            loader.load("/non/esiste.xlsx")


class TestWordLoader:
    def test_load_paragraphs(self, tmp_path):
        """Carica un Word con paragrafi di testo."""
        file = tmp_path / "test.docx"

        docx = DocxDocument()
        docx.add_paragraph("Primo paragrafo del documento.")
        docx.add_paragraph("Secondo paragrafo con più contenuto.")
        docx.save(str(file))

        loader = WordLoader()
        docs = loader.load(str(file))

        assert len(docs) == 1
        assert docs[0].doc_type == "docx"
        assert "Primo paragrafo" in docs[0].text
        assert "Secondo paragrafo" in docs[0].text

    def test_load_with_table(self, tmp_path):
        """Carica un Word con una tabella."""
        file = tmp_path / "table.docx"

        docx = DocxDocument()
        docx.add_paragraph("Introduzione al report.")

        table = docx.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Nome"
        table.cell(0, 1).text = "Valore"
        table.cell(1, 0).text = "Fatturato"
        table.cell(1, 1).text = "150000"
        table.cell(2, 0).text = "Costi"
        table.cell(2, 1).text = "80000"
        docx.save(str(file))

        loader = WordLoader()
        docs = loader.load(str(file))

        assert len(docs) == 1
        assert "Introduzione" in docs[0].text
        assert "Fatturato" in docs[0].text

    def test_file_not_found(self):
        loader = WordLoader()
        with pytest.raises(FileNotFoundError):
            loader.load("/non/esiste.docx")


class TestCSVLoader:
    def test_load_comma_separated(self, tmp_path):
        """Carica un CSV con virgola."""
        file = tmp_path / "data.csv"
        file.write_text(
            "Nome,Quantità,Prezzo\nUva,500,1.20\nPomodoro,300,0.80",
            encoding="utf-8",
        )

        loader = CSVLoader()
        docs = loader.load(str(file))

        assert len(docs) == 1
        assert "Nome: Uva" in docs[0].text
        assert "Prezzo: 1.20" in docs[0].text

    def test_load_semicolon_separated(self, tmp_path):
        """Carica un CSV con punto e virgola (formato italiano)."""
        file = tmp_path / "dati.csv"
        file.write_text(
            "Fornitore;Prodotto;Totale\nFerrara;Uva;1500\nRossi;Pomodoro;800",
            encoding="utf-8",
        )

        loader = CSVLoader()
        docs = loader.load(str(file))

        assert len(docs) == 1
        assert "Fornitore: Ferrara" in docs[0].text

    def test_file_not_found(self):
        loader = CSVLoader()
        with pytest.raises(FileNotFoundError):
            loader.load("/non/esiste.csv")


class TestLoaderRegistry:
    def test_get_pdf_loader(self):
        loader = get_loader("documento.pdf")
        assert isinstance(loader, PDFLoader)

    def test_get_excel_loader(self):
        loader = get_loader("bilancio.xlsx")
        assert isinstance(loader, ExcelLoader)

    def test_get_word_loader(self):
        loader = get_loader("contratto.docx")
        assert isinstance(loader, WordLoader)

    def test_get_csv_loader(self):
        loader = get_loader("dati.csv")
        assert isinstance(loader, CSVLoader)

    def test_get_txt_loader(self):
        loader = get_loader("note.txt")
        assert isinstance(loader, TXTLoader)

    def test_unsupported_format(self):
        with pytest.raises(ValueError, match="non supportato"):
            get_loader("immagine.png")

    def test_case_insensitive(self):
        """L'estensione è case-insensitive."""
        loader = get_loader("FILE.PDF")
        assert isinstance(loader, PDFLoader)

    def test_supported_extensions_complete(self):
        exts = supported_extensions()
        assert ".pdf" in exts
        assert ".xlsx" in exts
        assert ".docx" in exts
        assert ".csv" in exts
        assert ".txt" in exts
